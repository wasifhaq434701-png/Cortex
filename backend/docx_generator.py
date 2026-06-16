"""Phase 7 — server-side document export.

markdown_to_docx_bytes : python-docx  -> .docx bytes
markdown_to_pdf_bytes  : ReportLab     -> .pdf bytes

A lightweight Markdown subset is supported (headings, bold/italic inline,
bullet/numbered lists, blockquotes, code fences, paragraphs) — enough for the
synthesized research reports and edited documents. Both bundle cleanly into the
PyInstaller spec so export works offline inside the packaged app.
"""
import io
import re


# ---- shared tiny markdown tokenizer ---------------------------------------
def _iter_blocks(md: str):
    """Yield (kind, payload) blocks: ('h', (level, text)), ('li', (ordered, text)),
    ('quote', text), ('code', text), ('p', text)."""
    lines = md.replace("\r\n", "\n").split("\n")
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        # code fence
        if stripped.startswith("```"):
            buf = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1
            yield ("code", "\n".join(buf)); continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            yield ("h", (len(m.group(1)), m.group(2).strip())); i += 1; continue
        if stripped.startswith(("- ", "* ", "+ ")):
            yield ("li", (False, stripped[2:].strip())); i += 1; continue
        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            yield ("li", (True, m.group(1).strip())); i += 1; continue
        if stripped.startswith(">"):
            yield ("quote", stripped.lstrip("> ").strip()); i += 1; continue
        yield ("p", stripped); i += 1


_INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[.+?\]\(.+?\))")

def _strip_inline(text: str) -> str:
    """Plain-text version of inline markdown (for engines without rich runs)."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", text)
    return text


# ---- DOCX (python-docx) ----------------------------------------------------
def markdown_to_docx_bytes(md: str, title: str = "Document") -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    def add_inline_runs(paragraph, text):
        pos = 0
        for m in _INLINE.finditer(text):
            if m.start() > pos:
                paragraph.add_run(text[pos:m.start()])
            tok = m.group(0)
            if tok.startswith("**"):
                paragraph.add_run(tok[2:-2]).bold = True
            elif tok.startswith("*"):
                paragraph.add_run(tok[1:-1]).italic = True
            elif tok.startswith("`"):
                r = paragraph.add_run(tok[1:-1]); r.font.name = "Courier New"
            elif tok.startswith("["):
                lm = re.match(r"\[(.+?)\]\((.+?)\)", tok)
                paragraph.add_run(lm.group(1) if lm else tok)
            pos = m.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])

    for kind, payload in _iter_blocks(md):
        if kind == "h":
            level, text = payload
            doc.add_heading(_strip_inline(text), level=min(level, 4))
        elif kind == "li":
            ordered, text = payload
            p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
            add_inline_runs(p, text)
        elif kind == "quote":
            p = doc.add_paragraph(style="Intense Quote")
            add_inline_runs(p, payload)
        elif kind == "code":
            p = doc.add_paragraph()
            r = p.add_run(payload); r.font.name = "Courier New"; r.font.size = Pt(9)
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, payload)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---- PDF (ReportLab) -------------------------------------------------------
def markdown_to_pdf_bytes(md: str, title: str = "Document") -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, Preformatted

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=20 * mm, rightMargin=20 * mm,
                            topMargin=18 * mm, bottomMargin=18 * mm, title=title)
    styles = getSampleStyleSheet()
    code_style = ParagraphStyle("code", parent=styles["Code"], fontSize=8, leading=10)

    def inline_html(text):
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
        text = re.sub(r"`(.+?)`", r"<font face='Courier'>\1</font>", text)
        text = re.sub(r"\[(.+?)\]\((.+?)\)", r"<link href='\2' color='blue'>\1</link>", text)
        return text

    flow = []
    if title:
        flow.append(Paragraph(title, styles["Title"]))
        flow.append(Spacer(1, 6))

    pending_list = []
    def flush_list():
        nonlocal pending_list
        if pending_list:
            flow.append(ListFlowable([ListItem(Paragraph(inline_html(t), styles["BodyText"]))
                                      for t in pending_list], bulletType="bullet"))
            pending_list = []

    for kind, payload in _iter_blocks(md):
        if kind != "li":
            flush_list()
        if kind == "h":
            level, text = payload
            style = styles["Heading%d" % min(level, 4)] if ("Heading%d" % min(level, 4)) in styles else styles["Heading2"]
            flow.append(Paragraph(inline_html(text), style))
        elif kind == "li":
            pending_list.append(payload[1])
        elif kind == "quote":
            flow.append(Paragraph(inline_html(payload), styles["Italic"]))
        elif kind == "code":
            flow.append(Preformatted(payload, code_style))
        else:
            flow.append(Paragraph(inline_html(payload), styles["BodyText"]))
        flow.append(Spacer(1, 4))
    flush_list()

    doc.build(flow)
    return buf.getvalue()
